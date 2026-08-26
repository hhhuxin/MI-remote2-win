from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = r"C:\Users\Administrator\Documents\ChatGPT\小米蓝牙遥控器\服务成果验收报告模板.docx"

def set_cell(cell, text="", bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r.font.size = Pt(size)
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def shade(cell, fill="EDEDED"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def margins(cell, top=90, start=100, bottom=90, end=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            tcMar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def set_row_height(row, cm):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)

def set_width(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    tcW.set(qn("w:w"), str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(1.0), Cm(1.0)
sec.left_margin, sec.right_margin = Cm(1.25), Cm(1.25)

normal = doc.styles["Normal"]
normal.font.name = "宋体"; normal.font.size = Pt(9)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Q/CISDI OM11.52—2026")
r.font.name = "Times New Roman"; r.font.size = Pt(9)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(7)
r = p.add_run("附录 B：服务成果验收报告/服务确认文件模板")
r.bold = True; r.font.name = "黑体"; r.font.size = Pt(12)
r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")

t = doc.add_table(rows=4, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False; t.style = "Table Grid"
widths = [4.0, 6.3, 4.0, 4.2]
for row in t.rows:
    for i, c in enumerate(row.cells): set_width(c, widths[i]); margins(c)
for i, h in enumerate([0.85, 0.85, 1.05, 1.15]): set_row_height(t.rows[i], h)
labels = [(0,0,"项目名称"),(0,2,"项目编号"),(1,0,"合同名称"),(1,2,"合同编号"),(2,0,"合同签订时间"),(2,2,"验收报告提交/服务\n完成时间")]
for rr,cc,txt in labels: set_cell(t.cell(rr,cc),txt,bold=True,size=9); shade(t.cell(rr,cc),"F3F3F3")
t.cell(3,1).merge(t.cell(3,3)); set_cell(t.cell(3,0),"合同约定服务内容",bold=True,size=9); shade(t.cell(3,0),"F3F3F3")

mt = doc.add_table(rows=4, cols=6)
mt.alignment = WD_TABLE_ALIGNMENT.CENTER; mt.autofit = False; mt.style = "Table Grid"
mwidths = [2.6, 3.0, 4.4, 4.0, 3.2, 1.3]
for row in mt.rows:
    for i,c in enumerate(row.cells): set_width(c,mwidths[i]); margins(c,80,90,80,90)
for i,h in enumerate([0.85,2.9,1.7,2.0]): set_row_height(mt.rows[i],h)
headers = ["服务成果","提交时间","工作量认定","质量认定","交付进度认定","备注"]
for i,x in enumerate(headers): set_cell(mt.cell(0,i),x,bold=True,size=9,align=WD_ALIGN_PARAGRAPH.CENTER); shade(mt.cell(0,i),"E8E8E8")
set_cell(mt.cell(1,0),"成果 1")
set_cell(mt.cell(1,2),"xx%完成，小于 100%\n的须详细说明情况。",size=8.5)
set_cell(mt.cell(1,3),"优良/合格/让步接\n收/不合格，如为后\n两类须详细说明情\n况。",size=8.3)
set_cell(mt.cell(1,4),"提前/按期/延\n后，如为“延\n后”须详细说\n明情况。",size=8.3)
set_cell(mt.cell(2,0),"……")

def add_sign_block(title, roles=False):
    table = doc.add_table(rows=1, cols=1); table.style = "Table Grid"; table.autofit = False
    set_width(table.cell(0,0),18.5); margins(table.cell(0,0),100,130,80,130); set_row_height(table.rows[0],1.85)
    text = title
    if roles: text += "工程师、项目设计经理、项目经理："
    set_cell(table.cell(0,0), text + "\n\n                         签字：                 时间：          年     月     日", size=8.5)

add_sign_block("需求单位部门主管或主任", True)
add_sign_block("需求部门负责人、项目主管领导：")
add_sign_block("需求单位负责人：")

note = doc.add_table(rows=1, cols=1); note.style = "Table Grid"; note.autofit = False
set_width(note.cell(0,0),18.5); margins(note.cell(0,0),100,130,80,130); set_row_height(note.rows[0],2.2)
set_cell(note.cell(0,0),"备注：",bold=True,size=9)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(0)
p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.right_indent = Cm(0.5)
p.paragraph_format.line_spacing = 1.05
r = p.add_run("说明：需求单位负责合同执行，负责服务任务委托、服务工作量的认定、交付成果验收、服务质量评价等；负责按照合同约定向采购工程师提供合同服务结果验收报告及交付成果等作为支付有效依据文件。")
r.font.name = "宋体"; r.font.size = Pt(7.5); r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")

doc.save(OUT)
print(OUT)
